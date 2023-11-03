import spacy
from spacy.matcher import Matcher
from spacy.matcher import PhraseMatcher
import os
import glob
from PyPDF2 import PdfReader
import json
import re
import nltk
from nltk.corpus import stopwords
from docx import Document

# nltk.download("stopwords")

class ResumeParser:

    def __init__(self, pdf_folder, file_name):
        self.pdf_folder = pdf_folder
        self.nlp = spacy.load("en_core_web_lg")
        self.matcher = Matcher(self.nlp.vocab)
        skills_file="LINKEDIN_SKILLS_ORIGINAL.txt"
        self.skills_keywords = self.load_skills_keywords(skills_file)
        self.file_name = file_name

    def extract_text_from_pdf(self, pdf_file):
        with open(pdf_file, "rb") as pdf_file:
            pdf_reader = PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text

    def extract_text_from_docx(self, docx_file):
        doc = Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text
        return text

    def extract_date_of_birth(self, resume_text):
        doc = self.nlp(resume_text)
        dob_pattern_numeric = r'(\d{1,2}[-/]\d{1,2}[-/]\d{4})'
        dob_numeric = None
        
        for match in re.finditer(dob_pattern_numeric, resume_text):
            dob_numeric = match.group(0)
            break  # Stop after the first match

        # Check for natural language date of birth like "1st June 2001"
        dob_natural_language = None

        for token in doc:
            if token.text.isnumeric() and token.text.isdigit() and token.i + 1 < len(doc) and doc[token.i + 1].text.lower() in ["st", "nd", "rd", "th"]:
                date_tokens = [token]
                for i in range(token.i + 1, min(token.i + 4, len(doc))):  # Capture up to 4 tokens
                    date_tokens.append(doc[i])

                dob_natural_language = " ".join(token.text for token in date_tokens)
                break  # Stop after the first match

        return dob_numeric, dob_natural_language

    def extract_education(self, resume_text):
        doc = self.nlp(resume_text)

        degree = "Unknown"
        start_date = "Unknown"
        end_date = "Unknown"
        institution_names = []

        # Define keywords for education entity recognition
        education_keywords = ["bachelor's", "master's", "phd", "diploma", "degree",
                              "b.tech","m.tech","engineering","mba","b.e","be","bca","mca",
                              "b.com","m.com","bcom","mcom","bachelors","masters","bachelor","master"
                              "education"]

        for token in doc:
            if token.text.lower() in education_keywords:
                # Collect words following the education keywords as the degree
                degree = " ".join([t.text for t in token.subtree if not t.is_punct])
                # Break after finding the first degree
                break

        for i, token in enumerate(doc):
            if re.match(r'\b\d{4}\b', token.text):
                # Check if the previous token is a possible start date indicator
                if i > 0 and re.match(r'\b(year|date|from|since)\b', doc[i - 1].text, re.IGNORECASE):
                    start_date = token.text
                # Check if the next token is a possible end date indicator
                elif i < len(doc) - 1 and re.match(r'\b(year|date|to|until)\b', doc[i + 1].text, re.IGNORECASE):
                    end_date = token.text
        # Extract educational institution names
        for ent in doc.ents:
            if ent.label_ == "ORG":
                org_name = ent.text
                if "college" in org_name.lower() or "university" in org_name.lower() or "school" in org_name.lower() or "institute" in org_name.lower() or "institution" in org_name.lower():
                    institution_names.append(org_name)

        if not institution_names:
            institution_names.append("Unknown")

        return degree, start_date, end_date, institution_names

    def extract_additional_info(self, resume_text):
        doc = self.nlp(resume_text)
    
        contact_match = re.search(r'\d{10}', resume_text)
        contact_number = contact_match.group() if contact_match else "Unknown"
        country_codes=["+91","+92","+64"]
        
        if country_codes:
                # Try formatting the contact number with each country code
                for code in country_codes:
                    formatted_contact = f"{code}{contact_number}"
                    # Check if the formatted contact number is valid
                    # You may want to add additional validation as needed
                    if len(formatted_contact) <= 15:
                        contact_number = formatted_contact
                        break

        email_match = re.search(r'\S+@\S+', resume_text)
        email_address = email_match.group() if email_match else "Unknown"

        return contact_number, email_address

    def extract_name(self, resume_text):
        doc = self.nlp(resume_text)

        # Extract person's name using named entity recognition
        person_name = "Unknown"
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                person_name = ent.text
                break

        name_parts = person_name.split()
        if len(name_parts) == 2:
            first_name, last_name = name_parts
        else:
            first_name = name_parts[0]
            last_name = ""

        return first_name, last_name

    def extract_work_experience(self, resume_text):
        doc = self.nlp(resume_text)
        organizations = []
        current_company = None

        for sent in doc.sents:
            for ent in sent.ents:
                if ent.label_ == "ORG":
                    org_text = ent.text
                    # Check if "Present" is in the current sentence or its context
                    if "Present" in sent.text or "Present" in sent.text.lower():
                        current_company = org_text
                    if "Current" in sent.text or "Current" in sent.text.lower():
                        current_company = org_text
                    else:
                        organizations.append(org_text)

        return current_company, organizations

    # def extract_company(self, resume_text):
    #     custom_nlp = spacy.load(os.path.join("company_working","model"))
    #     doc = self.nlp(resume_text)
    #     co = []

    #     co = [ent.text.replace("\n", " ") for ent in list(doc.ents)]
    #     return list(dict.fromkeys(co).keys())


    def extract_urls(self, resume_text):
        # Define a regular expression pattern to match URLs with various formats
        # url_pattern = r'https?://\S+|www\.\S+|\S+\.\S+\.\S+'
        url_pattern = r'https?://\S+|www\.\S+'
        # Find all URLs in the resume text
        urls = re.findall(url_pattern, resume_text)

        return urls
        
    def extract_languages(self, resume_text):
        doc = self.nlp(resume_text)
        languages = []

        for token in doc:
            # Check if the token is tagged as a language (LANGUAGE entity in spaCy's NER)
            if token.ent_type_ == "LANGUAGE":
                languages.append(token.text)

        return languages

    def extract_addresses(self, resume_text):
        doc = self.nlp(resume_text)
        addresses = []

        for ent in doc.ents:
            if ent.label_ in ["GPE", "LOC"]:
                # Append recognized geographical entities (places) to the addresses
                addresses.append(ent.text)

        for token in doc:
            if token.text.lower() in ["apartment", "block", "lane", "road", "nagar", "street", "colony", "sector", "area",
                                     "room", "flat", "residence", "stay", "society", "plot", "house", "st."]:
                address = token.text
                # Collect words following the address term until a comma, period, or newline
                for next_token in token.rights:
                    if next_token.is_punct or next_token.is_space:
                        break
                    address += " " + next_token.text

                addresses.append(address)

            # Extract pincode using regular expression
            pincode_match = re.search(r'\b\d{6}\b', token.text)
            if pincode_match:
                addresses.append(pincode_match.group())

        return addresses

    def extract_designation(self, resume_text):
        # Load the spaCy model
        nlp = spacy.load("en_core_web_sm")
        designation_file="titles_combined.txt"
        # Read job designations from a file
        with open(designation_file, "r") as file:
            job_designations = [line.strip().lower() for line in file]

        # Initialize a PhraseMatcher
        matcher = PhraseMatcher(nlp.vocab)
        patterns = [nlp.make_doc(designation) for designation in job_designations]
        matcher.add("JobTitle", None, *patterns)

        # Process the resume text
        doc = nlp(resume_text.lower())

        # Use the PhraseMatcher to extract job designations
        filter_designation = [doc[start:end].text for match_id, start, end in matcher(doc)]

        return filter_designation

    def load_skills_keywords(self, skills_file):
        with open(skills_file, "r", encoding='utf-8') as file:
            skills = [line.strip().lower() for line in file if len(line.strip()) >= 4]
        return skills

    def extract_skills(self, resume_text):
        doc = self.nlp(resume_text)
        stop_words = set(stopwords.words("english"))
        skills = []

        for keyword in self.skills_keywords:
            if keyword.lower() in resume_text.lower() not in stop_words:
                skills.append(keyword)

        return list(set(skills))

    def extract_total_experience(self, resume_text):
        def correct_year(result):
            if len(result) < 2:
                if int(result) > int(str(date.today().year)[-2:]):
                    result = str(int(str(date.today().year)[:-2]) - 1) + result
                else:
                    result = str(date.today().year)[:-2] + result
            return result

        def calculate_experience(text):
            experience = 0
            start_month = -1
            start_year = -1
            end_month = -1
            end_year = -1

            not_alpha_numeric = r'[^a-zA-Z\d]'
            number = r'(\d{2})'

            months_num = r'(01)|(02)|(03)|(04)|(05)|(06)|(07)|(08)|(09)|(10)|(11)|(12)'
            months_short = r'(jan)|(feb)|(mar)|(apr)|(may)|(jun)|(jul)|(aug)|(sep)|(oct)|(nov)|(dec)'
            months_long = r'(january)|(february)|(march)|(april)|(may)|(june)|(july)|(august)|(september)|(october)|(november)|(december)'
            month = r'(' + months_num + r'|' + months_short + r'|' + months_long + r')'
            regex_year = r'((20|19)(\d{2})|(\d{2}))'
            year = regex_year
            start_date = month + not_alpha_numeric + r"?" + year

            end_date = r'((' + number + r'?' + not_alpha_numeric + r"?" + number + not_alpha_numeric + r"?" + year + r')|(present|current|till date|today))'
            longer_year = r"((20|19)(\d{2}))"
            year_range = longer_year + r"(" + not_alpha_numeric + r"{1,4}|(\s*to\s*))" + r'(' + longer_year + r'|(present|current|till date|today))'
            date_range = r"(" + start_date + r"(" + not_alpha_numeric + r"{1,4}|(\s*to\s*))" + end_date + r")|(" + year_range + r")"

            regular_expression = re.compile(date_range, re.IGNORECASE)
            regex_result = re.search(regular_expression, text)

            while regex_result:
                try:
                    date_range = regex_result.group()
                    try:
                        year_range_find = re.compile(year_range, re.IGNORECASE)
                        year_range_find = re.search(year_range_find, date_range)
                        replace = re.compile(
                            r"((\s*to\s*)|" + not_alpha_numeric + r"{1,4})", re.IGNORECASE)
                        replace = re.search(replace, year_range_find.group().strip())
                        start_year_result, end_year_result = year_range_find.group(
                        ).strip().split(replace.group())
                        start_year_result = int(correct_year(start_year_result))
                        if (end_year_result.lower().find('present') != -1 or end_year_result.lower().find('current') != -1 or
                                end_year_result.lower().find('till date') != -1 or end_year_result.lower().find('today') != -1):
                            end_month = date.today().month  # current month
                            end_year_result = date.today().year  # current year
                        else:
                            end_year_result = int(correct_year(end_year_result))
                    except Exception as e:
                        start_date_find = re.compile(start_date, re.IGNORECASE)
                        start_date_find = re.search(start_date_find, date_range)
                        non_alpha = re.compile(not_alpha_numeric, re.IGNORECASE)
                        non_alpha_find = re.search(non_alpha, start_date_find.group().strip())
                        replace = re.compile(
                            start_date + r"(" + not_alpha_numeric + r"{1,4}|(\s*to\s*))", re.IGNORECASE)
                        replace = re.search(replace, date_range)
                        date_range = date_range[replace.end():]
                        start_year_result = start_date_find.group().strip().split(
                            non_alpha_find.group())[-1]
                        start_year_result = int(correct_year(start_year_result))
                        if date_range.lower().find('present') != -1 or date_range.lower().find('current') != -1:
                            end_month = date.today().month  # current month
                            end_year_result = date.today().year  # current year
                        else:
                            end_date_find = re.compile(end_date, re.IGNORECASE)
                            end_date_find = re.search(end_date_find, date_range)
                            end_year_result = end_date_find.group().strip().split(
                                non_alpha_find.group())[-1]
                            try:
                                end_year_result = int(correct_year(end_year_result))
                            except Exception as e:
                                end_year_result = int(
                                    re.search("\d+", correct_year(end_year_result)).group())
                    if (start_year == -1) or (start_year_result <= start_year):
                        start_year = start_year_result
                    if (end_year == -1) or (end_year_result >= end_year):
                        end_year = end_year_result
                    text = text[regex_result.end():].strip()
                    regex_result = re.search(regular_expression, text)
                except Exception as e:
                    text = text[regex_result.end():].strip()
                    regex_result = re.search(regular_expression, text)

            return end_year - start_year

        total_experience = calculate_experience(resume_text)
        return total_experience
        
    def extract_certifications_and_courses(self, resume_text):
        # Load the English language model for spaCy
        nlp = spacy.load("en_core_web_sm")
        doc = self.nlp(resume_text)
        certifications = []
        
        # Iterate through the entities recognized by spaCy
        for ent in doc.ents:
            if ent.label_ == "CERTIFICATION":
                certifications.append(ent.text)

        # Use regular expressions to further extract certifications and courses
        certification_patterns = r"certification|certified|credential|licensed|accreditation|course|training|workshop|seminar|class|bootcamp"

        for sent in doc.sents:
            sentence = sent.text.lower()
            if re.search(certification_patterns, sentence):
                certifications.append(sent.text)

        # Deduplicate and return the extracted certifications and courses
        certifications = list(set(certifications))

        return certifications

    def extract_projects(self, resume_text):
        # Load the English language model for spaCy
        nlp = spacy.load("en_core_web_sm")

        # Process the resume text
        doc =self.nlp(resume_text)

        # Initialize a list to store projects
        projects = []

        # Define regular expressions to identify project-related terms
        project_patterns = r"project"

        # Iterate through the text and look for sentences mentioning projects
        for sent in doc.sents:
            sentence = sent.text.lower()
            if re.search(project_patterns, sentence):
                projects.append(sent.text)

        # Deduplicate and return the extracted projects
        projects = list(set(projects))

        return projects

    def process_resumes(self):
        pdf_files = glob.glob(os.path.join(self.pdf_folder, "*.pdf"))
        docx_files = glob.glob(os.path.join(self.pdf_folder, "*.docx"))
        parsed_resume = []

        for pdf_file in pdf_files:
            resume_text = self.extract_text_from_pdf(pdf_file)
            filter_designation=self.extract_designation(resume_text)
            filter_designation = list(dict.fromkeys(filter_designation).keys())
            job_role= self.extract_keywords(resume_text)
            degree, start_date, end_date, institution_names = self.extract_education(resume_text)
            contact_number, email_address = self.extract_additional_info(resume_text)
            current_company, organizations = self.extract_work_experience(resume_text)
            dob_numeric, dob_natural_language = self.extract_date_of_birth(resume_text)
            addresses=self.extract_addresses(resume_text)
            urls= self.extract_urls(resume_text)
            first_name, last_name =self.extract_name(resume_text)
            skills = self.extract_skills(resume_text)
            total_experience =self.extract_total_experience(resume_text)
            certifications= self.extract_certifications_and_courses(resume_text)
            projects = self.extract_projects(resume_text)

            parsed_resume.append({
                "data": {
                "firstName": first_name,
                "lastName":last_name,
                "fullName":"",
                "address":addresses,
                "current_designation":filter_designation,
                "experience":total_experience,
                "salary_expectation": "",
                "notice_period": "",
                "current_company": current_company,
                "email":email_address,
                "dob": dob_natural_language if dob_natural_language else dob_numeric,
                "mobile": contact_number,
                "work_mobile": "",
                "description": "",
                "resume": "",
                "key_skills": "",
                "technical_skills": [
                {
                    "skill_id": "",
                    "experience": "",
                },
                {
                    "skill_id": "",
                    "experience": "",
                },
                {
                    "skill_id": "",
                    "experience": 3,
                }
                ],
                "photo": "",
                "currentCTC": "",
                "currentLocation": "",
                "buyout_option": "",
                "location_preferred": "",
                "other_details": "",
                "educations": [
                {
                    "education": institution_names,
                    "course": degree,
                    "speciallization": "",
                    "achievement": "",
                    "university": "",
                    "courseType": "",
                    "passingOutYear": end_date,
                    "gradingSystem": ""
                },
                ],
                "employments": [
                {
                    "isCurrentEmloyment": 1, #yes
                    "employmentType": "",
                    "companyName": current_company,
                    "designation": "",
                    "joinDate": "",
                    "endDate": "",
                    "salary": "",
                    "jobProfile": ""
                },
                # {
                #     "isCurrentEmloyment": 2, #No
                #     "employmentType": "",
                #     "companyName": "",
                #     "designation": "",
                #     "joinDate": "",
                #     "endDate": "",
                #     "salary": "",
                #     "jobProfile": ""
                # }
                ],
                "projects": [
                {
                    "title": "",
                    "client": "",
                    "project_status": "",
                    "worked_from": "",
                    "worked_to": "",
                    "description": projects,
                },
                ]
            }})

        for docx_file in docx_files:
            resume_text = self.extract_text_from_docx(docx_file)
            filter_designation=self.extract_designation(resume_text)
            filter_designation = list(dict.fromkeys(filter_designation).keys())
            degree, start_date, end_date, institution_names = self.extract_education(resume_text)
            contact_number, email_address = self.extract_additional_info(resume_text)
            current_company, organizations = self.extract_work_experience(resume_text)
            dob_numeric, dob_natural_language = self.extract_date_of_birth(resume_text)
            addresses=self.extract_addresses(resume_text)
            urls= self.extract_urls(resume_text)
            first_name, last_name =self.extract_name(resume_text)
            skills = self.extract_skills(resume_text)
            total_experience =self.extract_total_experience(resume_text)
            certifications = self.extract_certifications_and_courses(resume_text)
            projects = self.extract_projects(resume_text)


            parsed_resume.append({
                "data": {
                "firstName": first_name,
                "lastName":last_name,
                "fullName":"",
                "address":addresses,
                "current_designation":filter_designation,
                "experience":total_experience,
                "salary_expectation": "",
                "notice_period": "",
                "current_company": current_company,
                "email":email_address,
                "dob": dob_natural_language if dob_natural_language else dob_numeric,
                "mobile": contact_number,
                "work_mobile": "",
                "description": "",
                "resume": "",
                "key_skills": "",
                "technical_skills": [
                {
                    "skill_id": "",
                    "experience": "",
                },
                {
                    "skill_id": "",
                    "experience": "",
                },
                {
                    "skill_id": "",
                    "experience": 3,
                }
                ],
                "photo": "",
                "currentCTC": "",
                "currentLocation": "",
                "buyout_option": "",
                "location_preferred": "",
                "other_details": "",
                "educations": [
                {
                    "education": institution_names,
                    "course": degree,
                    "speciallization": "",
                    "achievement": "",
                    "university": "",
                    "courseType": "",
                    "passingOutYear": end_date,
                    "gradingSystem": ""
                },
                ],
                "employments": [
                {
                    "isCurrentEmloyment": 1, #yes
                    "employmentType": "",
                    "companyName": current_company,
                    "designation": "",
                    "joinDate": "",
                    "endDate": "",
                    "salary": "",
                    "jobProfile": ""
                },
                # {
                #     "isCurrentEmloyment": 2, #No
                #     "employmentType": "",
                #     "companyName": "",
                #     "designation": "",
                #     "joinDate": "",
                #     "endDate": "",
                #     "salary": "",
                #     "jobProfile": ""
                # }
                ],
                "projects": [
                {
                    "title": "",
                    "client": "",
                    "project_status": "",
                    "worked_from": "",
                    "worked_to": "",
                    "description": projects,
                },
                ]
            }})

        return parsed_resume

    def process_resume(self):
        # pdf_file = glob.glob(os.path.join(self.pdf_folder, self.file_name)[0])
        parsed_resume = []

        # PDF
        if '.pdf' in self.file_name:
            pdf_file = os.path.join(self.pdf_folder, self.file_name)

            resume_text = self.extract_text_from_pdf(pdf_file)
            filter_designation=self.extract_designation(resume_text)
            filter_designation = list(dict.fromkeys(filter_designation).keys())
            degree, start_date, end_date, institution_names = self.extract_education(resume_text)
            contact_number, email_address = self.extract_additional_info(resume_text)
            current_company, organizations = self.extract_work_experience(resume_text)
            dob_numeric, dob_natural_language = self.extract_date_of_birth(resume_text)
            addresses=self.extract_addresses(resume_text)
            urls= self.extract_urls(resume_text)
            first_name, last_name =self.extract_name(resume_text)
            skills = self.extract_skills(resume_text)
            total_experience =self.extract_total_experience(resume_text)
            certifications= self.extract_certifications_and_courses(resume_text)
            projects = self.extract_projects(resume_text)
            languages = self.extract_languages(resume_text)
            # co = self.extract_company(resume_text)

            parsed_resume.append({
                "data": {
                "firstName": first_name,
                "lastName":last_name,
                "fullName":"",
                "address":addresses,
                "current_designation":filter_designation,
                "experience":total_experience,
                "salary_expectation": "",
                "notice_period": "",
                "current_company": current_company,
                "email":email_address,
                "dob": dob_natural_language if dob_natural_language else dob_numeric,
                "mobile": contact_number,
                "work_mobile": "",
                "description": "",
                "resume": "",
                "key_skills": "",
                "technical_skills": [
                {
                    "skill_id": "",
                    "experience": "",
                },
                {
                    "skill_id": "",
                    "experience": "",
                },
                {
                    "skill_id": "",
                    "experience": 3,
                }
                ],
                "photo": "",
                "currentCTC": "",
                "currentLocation": "",
                "buyout_option": "",
                "location_preferred": "",
                "other_details": "",
                "educations": [
                {
                    "education": institution_names,
                    "course": degree,
                    "speciallization": "",
                    "achievement": "",
                    "university": "",
                    "courseType": "",
                    "passingOutYear": end_date,
                    "gradingSystem": ""
                },
                ],
                "employments": [
                {
                    "isCurrentEmloyment": 1, #yes
                    "employmentType": "",
                    "companyName": current_company,
                    "designation": "",
                    "joinDate": "",
                    "endDate": "",
                    "salary": "",
                    "jobProfile": ""
                },
                {
                    "isCurrentEmloyment": 2, #No
                    "employmentType": "",
                    "companyName": "",
                    "designation": "",
                    "joinDate": "",
                    "endDate": "",
                    "salary": "",
                    "jobProfile": ""
                }
                ],
                "projects": [
                {
                    "title": "",
                    "client": "",
                    "project_status": "",
                    "worked_from": "",
                    "worked_to": "",
                    "description": projects,
                },
                ]
            }})

            # parsed_resume.append({
            #     "File": os.path.basename(pdf_file),
            #     "Designation": filter_designation,
            #     "First Name": first_name,
            #     "Last_Name":last_name,
            #     "DOB": dob_natural_language if dob_natural_language else dob_numeric,
            #     "URLs":urls,
            #     # "Company":co,
            #     "Languages":languages,
            #     "Address":addresses,
            #     "Current Organization":current_company,
            #     "Last Organization": organizations,
            #     "Contact Number": contact_number,
            #     "Email Address": email_address,
            #     "Total Years of Experience": total_experience,
            #     "Education":{"Degree": degree,
            #                   "Institution Names": institution_names,
            #                   "Start date": start_date,
            #                   "End date": end_date,
            #                   },
            #     "Projects":projects,
            #     "Certifications & Courses":certifications,
            #     "Skills": skills,
            # })

        # DOCX
        if '.doc' in self.file_name:
            docx_file = os.path.join(self.pdf_folder, self.file_name)
            resume_text = self.extract_text_from_docx(docx_file)
            filter_designation=self.extract_designation(resume_text)
            filter_designation = list(dict.fromkeys(filter_designation).keys())
            degree, start_date, end_date, institution_names = self.extract_education(resume_text)
            contact_number, email_address = self.extract_additional_info(resume_text)
            current_company, organizations = self.extract_work_experience(resume_text)
            dob_numeric, dob_natural_language = self.extract_date_of_birth(resume_text)
            addresses=self.extract_addresses(resume_text)
            urls= self.extract_urls(resume_text)
            first_name, last_name =self.extract_name(resume_text)
            skills = self.extract_skills(resume_text)
            total_experience =self.extract_total_experience(resume_text)
            certifications = self.extract_certifications_and_courses(resume_text)
            projects = self.extract_projects(resume_text)
            languages = self.extract_languages(resume_text)
            # co = self.extract_company(resume_text)

            parsed_resume.append({
                "data": {
                "firstName": first_name,
                "lastName":last_name,
                "fullName":"",
                "address":addresses,
                "current_designation":filter_designation,
                "experience":total_experience,
                "salary_expectation": "",
                "notice_period": "",
                "current_company": current_company,
                "email":email_address,
                "dob": dob_natural_language if dob_natural_language else dob_numeric,
                "mobile": contact_number,
                "work_mobile": "",
                "description": "",
                "resume": "",
                "key_skills": "",
                "technical_skills": [
                {
                    "skill_id": "",
                    "experience": "",
                },
                {
                    "skill_id": "",
                    "experience": "",
                },
                {
                    "skill_id": "",
                    "experience": 3,
                }
                ],
                "photo": "",
                "currentCTC": "",
                "currentLocation": "",
                "buyout_option": "",
                "location_preferred": "",
                "other_details": "",
                "educations": [
                {
                    "education": institution_names,
                    "course": degree,
                    "speciallization": "",
                    "achievement": "",
                    "university": "",
                    "courseType": "",
                    "passingOutYear": end_date,
                    "gradingSystem": ""
                },
                ],
                "employments": [
                {
                    "isCurrentEmloyment": 1, #yes
                    "employmentType": "",
                    "companyName": current_company,
                    "designation": "",
                    "joinDate": "",
                    "endDate": "",
                    "salary": "",
                    "jobProfile": ""
                },
                {
                    "isCurrentEmloyment": 2, #No
                    "employmentType": "",
                    "companyName": "",
                    "designation": "",
                    "joinDate": "",
                    "endDate": "",
                    "salary": "",
                    "jobProfile": ""
                }
                ],
                "projects": [
                {
                    "title": "",
                    "client": "",
                    "project_status": "",
                    "worked_from": "",
                    "worked_to": "",
                    "description": projects,
                },
                ]
            }})
            # parsed_resume.append({
            #     "File": os.path.basename(docx_file),
            #     "Designation": filter_designation,
            #     "First Name": first_name,
            #     "Last_Name":last_name,
            #     "DOB": dob_natural_language if dob_natural_language else dob_numeric,
            #     "URLs":urls,
            #     # "Company": co,
            #     "Languages":languages,
            #     "Address":addresses,
            #     "Current Organization":current_company,
            #     "Last Organization": organizations,
            #     "Contact Number": contact_number,
            #     "Email Address": email_address,
            #     "Total Years of Experience": total_experience,
            #     "Education":{"Degree": degree,
            #                   "Institution Names": institution_names,
            #                   "Start date": start_date,
            #                   "End date": end_date,
            #                   },
            #     "Projects":projects,
            #     "Certifications & Courses":certifications,
            #     "Skills": skills,
            # })

        return parsed_resume[0]

# if __name__ == "__main__":
# Root folder containing subfolders with PDFs and DOCX files
root_folder = r'ex'

# Output folder for JSON and text files
output_folder = r"output_folder"
os.makedirs(output_folder, exist_ok=True)

# resume_parser = ResumeParser(os.path.join('ex', 'Profiles I'), 'Goutham Korepu_SOC.pdf')
# parsed_resume = resume_parser.process_resume()
# print(parsed_resume)

def get_json(file, filename):
    resume_parser = ResumeParser(file, filename)
    parsed_resume = resume_parser.process_resume()
    return parsed_resume
    
# for subfolder in os.listdir(root_folder):
#     subfolder_path = os.path.join(root_folder, subfolder)
#     if os.path.isdir(subfolder_path):
#         resume_parser = ResumeParser(subfolder_path)
#         parsed_resume = resume_parser.process_resumes()

#         # Create a subfolder for each subfolder in the output directory
#         subfolder_output_folder = os.path.join(output_folder, subfolder)
#         os.makedirs(subfolder_output_folder, exist_ok=True)

#         for resume in parsed_resume:
#             output_text_file = os.path.join(subfolder_output_folder, f"{resume['File']}.txt")
#             with open(output_text_file, "w", encoding="utf-8") as text_file:
#                 text_file.write(f"File: {resume['File']}\n")
#                 text_file.write(f"Category (Job Role): {resume['Category (Job Role)']}\n")
#                 text_file.write(f"Person Name: {resume['Person Name']}\n")
#                 text_file.write(f"Contact Details: {resume['Contact Details']}\n")
#                 text_file.write(f"Total Years of Experience: {resume['Total Years of Experience']}\n")
#                 text_file.write(f"Education: Degree - {resume['Education']['Degree']}\n")
#                 text_file.write(f"Institution Names: {', '.join(resume['Education']['Institution Names'])}\n")
#                 text_file.write(f"Skills: {', '.join(resume['Skills'])}\n")

#             output_json_file = os.path.join(subfolder_output_folder, f"{os.path.splitext(resume['File'])[0]}.json")
#             with open(output_json_file, "w", encoding="utf-8") as json_file:
#                 json.dump(resume, json_file, indent=4)

# def parseDocument():
